# -*- coding: utf-8 -*-
# Document Processing Module for AI Document Analyzer
# Handles text extraction from PDF, Word, and text files

import PyPDF2
import docx
import io
from typing import Dict, List, Optional
import re

class DocumentProcessor:
    """
    Handles processing of various document formats including PDF, Word, and text files.
    Extracts text content and splits into manageable chunks for analysis.
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the document processor.
        
        Args:
            chunk_size (int): Size of text chunks in characters
            chunk_overlap (int): Overlap between chunks in characters
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_document(self, file, filename: str) -> Dict:
        """
        Process uploaded document and extract text content.
        
        Args:
            file: Uploaded file object
            filename (str): Name of the uploaded file
            
        Returns:
            Dict: Processed document information including text and chunks
        """
        try:
            # Determine file type and extract text
            if filename.lower().endswith('.pdf'):
                text = self._extract_pdf_text(file)
                file_type = "PDF"
            elif filename.lower().endswith(('.docx', '.doc')):
                text = self._extract_word_text(file)
                file_type = "Word Document"
            elif filename.lower().endswith('.txt'):
                text = self._extract_text_file(file)
                file_type = "Text File"
            else:
                raise ValueError(f"Unsupported file type: {filename}")
            
            # Clean and process text
            cleaned_text = self._clean_text(text)
            
            # Split into chunks
            chunks = self._split_into_chunks(cleaned_text)
            
            # Calculate statistics
            word_count = len(cleaned_text.split())
            char_count = len(cleaned_text)
            
            return {
                "filename": filename,
                "file_type": file_type,
                "text": cleaned_text,
                "chunks": chunks,
                "word_count": word_count,
                "character_count": char_count,
                "chunk_count": len(chunks),
                "success": True,
                "error": None
            }
            
        except Exception as e:
            return {
                "filename": filename,
                "file_type": "Unknown",
                "text": "",
                "chunks": [],
                "word_count": 0,
                "character_count": 0,
                "chunk_count": 0,
                "success": False,
                "error": str(e)
            }
    
    def _extract_pdf_text(self, file) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            # Reset file pointer
            file.seek(0)
            
            # Create PDF reader
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Extract text from all pages
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
        
        return text
    
    def _extract_word_text(self, file) -> str:
        """Extract text from Word document"""
        try:
            # Reset file pointer
            file.seek(0)
            
            # Read Word document
            doc = docx.Document(file)
            
            # Extract text from all paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
        except Exception as e:
            raise Exception(f"Error reading Word document: {str(e)}")
        
        return text
    
    def _extract_text_file(self, file) -> str:
        """Extract text from plain text file"""
        try:
            # Reset file pointer
            file.seek(0)
            
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    file.seek(0)
                    content = file.read()
                    if isinstance(content, bytes):
                        text = content.decode(encoding)
                    else:
                        text = content
                    return text
                except (UnicodeDecodeError, UnicodeError):
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            file.seek(0)
            content = file.read()
            if isinstance(content, bytes):
                text = content.decode('utf-8', errors='ignore')
            else:
                text = content
            return text
            
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s.,!?;:()\-\'\"]+', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _split_into_chunks(self, text: str) -> List[Dict]:
        """Split text into overlapping chunks for better processing"""
        if not text:
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Calculate end position
            end = start + self.chunk_size
            
            # If this isn't the last chunk, try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                for i in range(end, max(start + self.chunk_size - 100, start), -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
            
            # Extract chunk
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunks.append({
                    "index": len(chunks),
                    "text": chunk_text,
                    "start_pos": start,
                    "end_pos": end,
                    "word_count": len(chunk_text.split())
                })
            
            # Move to next chunk with overlap
            start = max(start + self.chunk_size - self.chunk_overlap, end)
            
            # Prevent infinite loop
            if start >= len(text):
                break
        
        return chunks
    
    def get_document_summary(self, doc_info: Dict) -> str:
        """Generate a brief summary of the processed document"""
        if not doc_info["success"]:
            return f"Error processing {doc_info['filename']}: {doc_info['error']}"
        
        return f"""
**Document Summary:**
- **File:** {doc_info['filename']}
- **Type:** {doc_info['file_type']}
- **Words:** {doc_info['word_count']:,}
- **Characters:** {doc_info['character_count']:,}
- **Chunks:** {doc_info['chunk_count']}
"""
    
    def search_chunks(self, doc_info: Dict, query: str, max_results: int = 3) -> List[Dict]:
        """
        Simple text search within document chunks.
        
        Args:
            doc_info (Dict): Processed document information
            query (str): Search query
            max_results (int): Maximum number of results to return
            
        Returns:
            List[Dict]: Matching chunks with relevance scores
        """
        if not doc_info["success"] or not query:
            return []
        
        query_words = query.lower().split()
        results = []
        
        for chunk in doc_info["chunks"]:
            chunk_text = chunk["text"].lower()
            
            # Calculate simple relevance score
            score = 0
            for word in query_words:
                score += chunk_text.count(word)
            
            if score > 0:
                results.append({
                    "chunk": chunk,
                    "score": score,
                    "relevance": min(score / len(query_words), 1.0)
                })
        
        # Sort by relevance and return top results
        results.sort(key=lambda x: x["score"], reverse=True)
        return results[:max_results]